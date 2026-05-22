"""
Generate expanded Capstone-I report for Go Bharat (Group 138).
Includes Agentic AI theory, single vs multi-agent analysis, and diagrams.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ASSETS = Path(__file__).resolve().parent / "report_assets"
OUT_PATH = Path(__file__).resolve().parent.parent / "Go_Bharat_Capstone_Report_Group138_LangGraph_v2.docx"

NAVY = RGBColor(0x1A, 0x36, 0x5D)
TEAL = RGBColor(0x0D, 0x5C, 0x63)
SAFFRON = RGBColor(0xE8, 0x6C, 0x00)
BODY = RGBColor(0x2C, 0x3E, 0x50)
MUTED = RGBColor(0x5D, 0x6D, 0x7E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FONT_BODY = "Calibri"
FONT_HEADING = "Georgia"
REPORT_DATE = "30 March 2026"
LIVE_URL = "https://travel-agent-using-langchain.onrender.com"


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E86C00")
    pBdr.append(bottom)
    pPr.append(pBdr)


def style_paragraph(p, *, size=11, align=None, space_after=8):
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.2
    if align is not None:
        p.alignment = align
    for run in p.runs:
        run.font.name = FONT_BODY
        run.font.size = Pt(size)
        run.font.color.rgb = BODY


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    style_paragraph(p)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    configs = {1: (16, NAVY, 18, 10), 2: (13, TEAL, 14, 8), 3: (11, TEAL, 10, 6)}
    size, color, before, after = configs.get(level, configs[3])
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_HEADING if level <= 2 else FONT_BODY
    run.font.size = Pt(size)
    run.font.color.rgb = color


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        style_paragraph(p)


def add_numbered(doc: Document, items: list[str]) -> None:
    """Word continues list numbers across the document — prefer add_numbered_fresh or add_bullets."""
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        style_paragraph(p)


def add_numbered_fresh(doc: Document, items: list[str]) -> None:
    """Numbered list 1, 2, 3… that does not continue from earlier lists (e.g. TOC)."""
    for i, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {item}")
        run.font.name = FONT_BODY
        run.font.size = Pt(11)
        run.font.color.rgb = BODY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.2


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.2) -> None:
    if not image_path.exists():
        add_body(doc, f"[Figure unavailable: {image_path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED
        r.font.name = FONT_BODY
    doc.add_paragraph()


def add_comparison_table(doc: Document, headers: list[str], rows: list[tuple]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "1A365D")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    doc.add_paragraph()


def build_document() -> Document:
    from report_diagrams import generate_all

    generate_all()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

    # ===== TITLE PAGE =====
    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t1.add_run("GO BHARAT")
    r.bold = True
    r.font.name = FONT_HEADING
    r.font.size = Pt(30)
    r.font.color.rgb = NAVY

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Smart Travel Planning Website for Digital Bharat")
    r2.font.name = FONT_HEADING
    r2.font.size = Pt(15)
    r2.font.color.rgb = SAFFRON
    r2.italic = True

    doc.add_paragraph()
    for line in [
        "Capstone-I Project Report",
        "Hybrid UG Programme in Computer Science & Data Analytics",
        "Indian Institute of Technology Patna, Bihta — 801106, India",
        f"Group No. 138  |  Date: {REPORT_DATE}",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    student = doc.add_paragraph()
    student.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = student.add_run("Submitted by\nRitika Rai\nRoll No. UA2503CDH475")
    sr.font.name = FONT_BODY
    sr.font.size = Pt(12)
    sr.font.color.rgb = BODY

    doc.add_page_break()

    # ===== TABLE OF CONTENTS (manual) =====
    add_heading(doc, "Table of Contents", 1)
    toc = [
        "Declaration",
        "Abstract",
        "List of Figures",
        "Team Contributions",
        "My Contribution",
        "1. Introduction",
        "2. Problem Statement",
        "3. Objectives",
        "4. Proposed Solution",
        "5. System Architecture",
        "6. System Features",
        "7. Understanding Agentic AI",
        "8. Single Agent vs Multi-Agent Design",
        "9. Agentic AI in Go Bharat",
        "10. Technologies Used",
        "11. Database Design",
        "12. Methodology",
        "13. Testing Approach",
        "14. Advantages and Limitations",
        "15. API Overview",
        "11. Database Design (Detailed)",
        "12–18. Methodology through References",
        "19. Deployment Architecture",
        "20. Layered Software Architecture",
        "21. API Request Sequence",
        "Appendices A–F",
    ]
    add_bullets(doc, toc)
    doc.add_page_break()

    # ===== DECLARATION =====
    add_heading(doc, "Declaration", 1)
    add_body(
        doc,
        "I hereby declare that this submission is my own work and that, to the best of my knowledge "
        "and belief, it contains no material previously published or written by another person, nor "
        "material which to a substantial extent has been accepted for the award of any other degree "
        "or diploma of the university or other institute of higher learning, except where due "
        "acknowledgement has been made in the text.",
    )
    add_body(doc, f"Date: {REPORT_DATE}\nName: Ritika Rai\nRoll number: UA2503CDH475\nGroup number: 138\n[Signature]")
    add_horizontal_rule(doc)

    # ===== ABSTRACT (expanded) =====
    add_heading(doc, "Abstract", 1)
    for para in [
        "Tourism is a major contributor to India's economy and to the personal growth of millions of "
        "travellers. Yet trip planning remains fragmented: weather, accommodation, dining, sightseeing, "
        "and budgeting are spread across dozens of websites and apps. Go Bharat (Group 138, IIT Patna) "
        "addresses this by building a production-style web application that unifies live travel data with "
        "agentic artificial intelligence.",
        "The platform accepts natural-language or structured travel requirements (destination, dates, budget, "
        "and preferences). A LangGraph-based multi-agent pipeline processes each request: a Parser Agent "
        "extracts entities and routes work to specialist agents that run in parallel—Weather, Hotel, "
        "Restaurants, Attractions, Itinerary, Budget, and Things to Carry—before a Merge step combines "
        "outputs into a single Summary Plan. Live APIs (Open-Meteo, TripAdvisor via RapidAPI, OpenTripMap, "
        "exchange rates) ground each agent. A Travel Chatbot supports follow-up Q&A on saved plans.",
        "Technically, the stack includes FastAPI, LangGraph, LangChain, Groq LLMs, Supabase (PostgreSQL), "
        "and deployment on Render. The design demonstrates why parallel specialist agents outperform one "
        "monolithic LLM for complex travel workflows.",
    ]:
        add_body(doc, para)
    add_horizontal_rule(doc)

    add_heading(doc, "List of Figures", 2)
    add_bullets(
        doc,
        [
            "Figure 1 — System architecture of Go Bharat",
            "Figure 2 — Comparison: single monolithic agent vs multi-agent design",
            "Figure 3 — LangGraph multi-agent travel planner flow",
            "Figure 4 — Dashboard UI concept",
            "Figure 5 — Database layout (travel_plans + JSONB)",
            "Figure 6 — Database CRUD operations",
            "Figure 7 — Sample travel_plans row (table view)",
            "Figure 8 — API sequence for POST /plan",
            "Figure 9 — Deployment architecture (Render + Supabase + Groq)",
            "Figure 10 — Layered application architecture",
        ],
    )
    doc.add_page_break()

    # ===== TEAM =====
    add_heading(doc, "Team Contributions", 1)
    team = [
        ("Nishant Rai", "Project concept, scope definition, and overall vision for the travel platform."),
        ("Atharva Rai", "Backend architecture: FastAPI, LangGraph multi-agent pipeline, API integrations, and deployment."),
        ("Pranav Rai", "Supporting development and learning full-stack components alongside the team."),
        ("Avni Rai", "Data collection and supporting project tasks during implementation."),
        ("Ritika Rai", "Frontend development (HTML, CSS, Jinja2 templates), UI/UX polish, presentation materials, and integration support."),
    ]
    add_comparison_table(doc, ["Member", "Role"], team)

    add_heading(doc, "My Contribution (Ritika Rai)", 2)
    for para in [
        "I focused primarily on the presentation layer of Go Bharat: designing and implementing server-rendered "
        "pages using HTML, CSS, and Jinja2 templates integrated with FastAPI. I worked to make the interface "
        "responsive, readable, and easy to navigate for users entering trip details and viewing generated plans "
        "on the dashboard.",
        "I collaborated on the project presentation and supported integration between the frontend forms and "
        "backend REST endpoints. Understanding how Supabase stores JSON travel plans helped me see how frontend "
        "actions (create plan, view history, chat) depend on persisted backend state.",
        "Working with Atharva on the LangGraph design taught me that AI in production is not one giant chatbot—it is "
        "a Parser plus parallel specialists (weather, hotels, food, attractions, itinerary, budget, packing) merged "
        "into one Summary Plan. I contributed to communicating this flow in our slides and documentation.",
        "Skills gained: responsive web layout, template-driven UIs, API-driven pages, teamwork, and exposure to "
        "LangChain-based agents without building the entire backend alone.",
    ]:
        add_body(doc, para)
    doc.add_page_break()

    # ===== 1 INTRODUCTION (expanded) =====
    add_heading(doc, "1. Introduction", 1)
    for para in [
        "Travel has become essential for education, leisure, business, and cultural exploration. The Government "
        "of India's National Digital Tourism Mission and related policy initiatives aim to make tourism information "
        "more accessible through technology. Despite this, individual travellers still spend hours comparing "
        "blogs, videos, maps, and booking sites before a single trip.",
        "Most existing platforms optimize for popular destinations and generic listings. They rarely combine "
        "live weather, structured hotel and restaurant data, attraction discovery, currency-aware budgeting, and "
        "conversational refinement in one coherent workflow. Students and families planning trips across India "
        "often need practical, budget-conscious guidance—not only landmark names.",
        "Go Bharat is our team's answer: a smart travel companion that produces a structured trip brief from "
        "real APIs, stores it for the user, and layers optional AI agents for weather narrative, itineraries, "
        "budget tables, and chat. The name reflects a digital platform aligned with the spirit of accessible "
        "travel planning for Bharat (India) and beyond.",
        "This report explains the motivation, technical stack, and—centrally—why we adopted an agentic, "
        "multi-agent AI architecture instead of relying on one all-purpose language model.",
    ]:
        add_body(doc, para)

    add_heading(doc, "2. Problem Statement", 1)
    add_body(doc, "Travellers today face several recurring difficulties:")
    add_bullets(
        doc,
        [
            "Fragmented information across weather portals, hotel sites, food blogs, and map applications.",
            "Difficulty comparing options when budget, dates, and travel style must all be respected together.",
            "Generic recommendations that ignore party size, season, or currency.",
            "No persistent plan object—users re-enter the same details on every new website.",
            "Chatbots that answer in isolation without access to the user's actual trip data.",
            "Risk of hallucination when a single model invents hotels or weather without API grounding.",
        ],
    )

    add_heading(doc, "3. Objectives", 1)
    add_bullets(
        doc,
        [
            "Build a unified travel planning web application with validated inputs (Pydantic schemas).",
            "Integrate live third-party APIs for weather, hotels, restaurants, attractions, and currency.",
            "Persist plans in Supabase for retrieval by email and plan identifier.",
            "Implement a LangGraph workflow with Parser Agent, seven parallel specialist agents, Merge, and Summary Plan.",
            "Ground each agent with live APIs (weather, hotels, food, attractions, budget).",
            "Offer a plan-aware Travel Chatbot for follow-up questions on stored plans.",
            "Deploy a demonstrable production URL for evaluators and users.",
            "Document why multi-agent design is superior to a single-agent approach for this domain.",
        ],
    )

    add_heading(doc, "4. Proposed Solution", 1)
    add_body(
        doc,
        "Go Bharat is implemented as a layered FastAPI application fronting a LangGraph orchestration graph. "
        "User travel requirements enter the graph at User Input; the Parser Agent understands the query, "
        "extracts entities (city, dates, budget, style), and triggers parallel execution of seven specialist "
        "agents. Each agent calls the relevant external APIs and/or LLM prompts. A Merge node collects partial "
        "results; the Summary Plan node produces the final comprehensive travel brief shown on the dashboard. "
        "Supabase persists the merged plan; Jinja2 renders the UI.",
    )
    add_body(
        doc,
        "FastAPI exposes REST endpoints (/plan, /itinerary, /budget, /chat) that invoke the LangGraph workflow "
        "or individual agent paths where appropriate. Pydantic validates dates (check-out after check-in) before "
        "the graph runs. The Travel Chatbot remains available for conversational refinement using the stored "
        "travel_details JSON after the Summary Plan is saved.",
    )

    add_heading(doc, "5. System Architecture", 1)
    add_body(doc, "Figure 1 shows how user requests flow through the web layer, orchestrator, agents, database, and LLM provider.")
    add_figure(doc, ASSETS / "fig1_system_architecture.png", "Figure 1: High-level system architecture of Go Bharat")
    add_bullets(
        doc,
        [
            "Presentation layer: login and dashboard pages (HTML/CSS/Jinja2).",
            "API layer: FastAPI router at /api/v1 with health, plan, plans, chat, itinerary, and budget routes.",
            "Service layer: LangGraph workflow, PlannerService/API tools, SupabaseService, specialist agents.",
            "Data layer: Supabase PostgreSQL table travel_plans (email_id, city, travel_details JSONB).",
            "External services: Open-Meteo, RapidAPI TripAdvisor16, OpenTripMap, Exchange Rate API, Groq.",
        ],
    )

    add_heading(doc, "6. System Features", 1)
    add_bullets(
        doc,
        [
            "Trip form: city, check-in/out, adults (1–10), budget amount/currency, travel style.",
            "Weather brief: current conditions plus historical expectation for the travel date.",
            "Hotel cards: ratings, price hints, and addresses from TripAdvisor search.",
            "Restaurant suggestions with ratings and price tags.",
            "Attraction cards from OpenTripMap radius search with categories.",
            "Currency conversion sentence for the user's budget.",
            "Dashboard listing of saved plans per email.",
            "Parallel LangGraph agents: weather, hotels, restaurants, attractions, itinerary, budget, packing list.",
            "Merged Summary Plan combining all agent outputs.",
            "Multi-turn Travel Chatbot referencing the saved plan.",
        ],
    )

    # ===== 7 AGENTIC AI THEORY (NEW - LONG) =====
    add_heading(doc, "7. Understanding Agentic AI", 1)

    add_heading(doc, "7.1 What is Agentic AI?", 2)
    for para in [
        "Artificial Intelligence (AI) in software can take many forms. Traditional rule-based systems follow "
        "fixed if–else logic. Machine learning models classify or predict from data. Large Language Models (LLMs) "
        "generate human-like text from prompts. Agentic AI goes further: it describes systems where an AI component "
        "is given a goal, context, and often tools, and can plan steps toward that goal rather than producing a "
        "single static reply.",
        "In travel technology, an agentic module might read live weather numbers, decide what matters for a "
        "traveller (heat, rain, packing), and write two short paragraphs—in a tone suitable for tourists. Another "
        "agent might receive hotel lists and produce a day-by-day schedule in table form. Each agent has a "
        "contract: inputs it expects, outputs it must produce, and boundaries on what it should not invent.",
        "Agentic AI does not mean the software is fully autonomous like a human travel agent. In Go Bharat, "
        "humans still choose the destination and budget; FastAPI still controls which agent runs and when; APIs "
        "still supply ground truth. The “agency” is specialized decision-making and language generation within "
        "those guardrails.",
    ]:
        add_body(doc, para)

    add_heading(doc, "7.2 Key characteristics of agentic systems", 2)
    add_bullets(
        doc,
        [
            "Goal-oriented: each agent has one task (e.g. hotels only, weather only, packing list only).",
            "Context-aware: prompts include structured data from databases and APIs, not only the user's last message.",
            "Composable: orchestrator code decides which agent to call and merges results into the UI.",
            "Inspectable: prompts and outputs can be logged and improved independently.",
            "Bounded: agents are discouraged from inventing facts outside provided lists (hotels, restaurants, etc.).",
        ],
    )

    add_heading(doc, "7.3 Agentic AI vs traditional chatbots", 2)
    add_comparison_table(
        doc,
        ["Aspect", "Traditional chatbot", "Agentic approach (Go Bharat)"],
        [
            ("Knowledge", "Fixed FAQ or keywords", "Live API data injected into prompts"),
            ("Scope", "One generic dialog", "LangGraph: Parser + parallel specialists + Merge"),
            ("Memory", "Often none or session-only", "Stored plan in Supabase + chat session"),
            ("Output", "Short text replies", "Tables, HTML briefs, structured markdown"),
            ("Maintenance", "Hard to update one bot for all tasks", "Update one agent without breaking others"),
        ],
    )

    add_heading(doc, "7.4 Role of LangGraph, LangChain, and Groq", 2)
    add_body(
        doc,
        "LangGraph models the travel planner as a directed graph: nodes are agents or tools; edges define "
        "control flow (sequential Parser → parallel fan-out → Merge → Summary). This matches Figure 3 and is "
        "the standard pattern for multi-agent systems that need both coordination and concurrency.",
    )
    add_body(
        doc,
        "LangChain provides prompt templates, output parsers, and runnable chains (prompt | llm | parser) inside "
        "each graph node. Each specialist wraps ChatGroq (model openai/gpt-oss-120b) with a task-specific template. "
        "Groq supplies fast inference for interactive requests. Temperature and max_tokens differ per agent—for "
        "example, Budget Agent allows longer tabular output than Weather Agent.",
    )

    doc.add_page_break()

    # ===== 8 SINGLE VS MULTI (NEW - LONG) =====
    add_heading(doc, "8. Single Agent vs Multi-Agent Design", 1)

    add_heading(doc, "8.1 The single-agent approach (and why we rejected it)", 2)
    add_body(
        doc,
        "A naive design would use one LLM call to “do everything”: parse the city, call APIs mentally, write "
        "weather, list hotels, build a 7-day itinerary, produce three budget tiers, and answer chat—all in one "
        "prompt. This monolithic pattern is common in demos but scales poorly in practice.",
    )
    add_figure(doc, ASSETS / "fig2_single_vs_multi_agent.png", "Figure 2: Single monolithic agent (left) vs multi-agent design used in Go Bharat (right)")

    add_heading(doc, "8.2 Why a single agent is a bad fit for travel planning", 2)
    add_body(doc, "The following problems appear when one model handles every travel task in one shot:")
    add_bullets(
        doc,
        [
            "Prompt overload: the model must juggle weather rules, table formats, budget math, and chat tone simultaneously, increasing confusion and omitted sections.",
            "Higher hallucination risk: without strict per-task context, the model may invent hotel names or prices not present in API results.",
            "Poor output structure: itineraries need markdown tables; weather needs short paragraphs; budgets need eight numbered sections—one prompt rarely enforces all formats reliably.",
            "Cost and latency: regenerating the entire mega-response for a small chat question wastes tokens and slows the user.",
            "No separation of concerns: a bug in budget formatting forces redeploying the same prompt that also handles weather.",
            "Harder testing: failures cannot be isolated to “itinerary only” or “budget only”.",
            "Concurrency limits: users cannot request only an itinerary update without re-running weather and hotels.",
            "Safety and moderation: travel advice mixed with open-ended chat in one thread is harder to audit.",
        ],
    )

    add_heading(doc, "8.3 Why multiple agents are better", 2)
    add_body(
        doc,
        "Multi-agent (or multi-chain) design assigns each concern to a module with its own prompt, inputs, and "
        "output schema. Go Bharat uses this pattern deliberately:",
    )
    add_bullets(
        doc,
        [
            "Specialization: each parallel agent has one prompt and one output schema (hotels vs weather vs packing).",
            "Parallel execution: seven agents run concurrently after parsing, reducing end-to-end latency.",
            "Grounding: Hotel and Restaurants agents use TripAdvisor API data; Attractions use OpenTripMap; Weather uses Open-Meteo.",
            "Independent iteration: Budget rules can change without redeploying the Hotel agent node.",
            "Clear mapping to Figure 3: Parser → parallel experts → Merge → Summary Plan.",
            "LangGraph adds explicit Merge and Summary nodes—cleaner than ad-hoc Python orchestration alone.",
            "Future scaling: new graph nodes (e.g. Visa Agent) plug in without rewriting the whole pipeline.",
        ],
    )

    add_comparison_table(
        doc,
        ["Criterion", "Single agent", "Multi-agent (Go Bharat)"],
        [
            ("Prompt complexity", "Very high", "Moderate per agent"),
            ("Hallucination control", "Weak", "Stronger (API lists in context)"),
            ("Response time for chat", "Slow (full replan)", "Fast (chatbot only)"),
            ("Maintainability", "Low", "High"),
            ("Academic clarity", "Hard to explain", "Clear modules for report/diagram"),
        ],
    )

    add_heading(doc, "8.4 When is a single agent acceptable?", 2)
    add_body(
        doc,
        "A single LLM may suffice for a toy demo or a narrow FAQ bot with ten static answers. For Go Bharat, "
        "the combination of live APIs, long structured outputs, persistent plans, and multi-turn chat exceeds "
        "what one prompt can reliably govern. Go Bharat uses LangGraph with a Parser, seven parallel specialists, "
        "Merge, and Summary Plan; the Travel Chatbot handles optional follow-up dialogue on stored plans.",
    )

    doc.add_page_break()

    # ===== 9 AGENTIC IN GO BHARAT (LangGraph) =====
    add_heading(doc, "9. Agentic AI in Go Bharat — LangGraph Architecture", 1)
    add_body(
        doc,
        "The core intelligence of Go Bharat follows the LangGraph multi-agent flow shown in Figure 3. "
        "This is the authoritative design for how user queries become a complete travel plan.",
    )
    add_figure(
        doc,
        ASSETS / "fig_langgraph_multi_agent.png",
        "Figure 3: LangGraph architecture — Travel Planner multi-agent flow (Parser → parallel agents → Merge → Summary Plan)",
        width=6.4,
    )

    add_heading(doc, "9.1 Pipeline stages", 2)
    add_numbered_fresh(
        doc,
        [
            "User Input — travel query or structured form (e.g. “Trip to Paris for 5 days”, or city + dates + budget).",
            "Parser Agent — understands intent, extracts entities (destination, duration, budget, interests), routes to experts.",
            "Parallel execution — seven specialist agents run concurrently (see Table below).",
            "Merge — aggregates structured outputs from all parallel branches.",
            "Summary Plan — combines insights into one comprehensive travel plan for the dashboard and Supabase.",
        ],
    )

    add_heading(doc, "9.2 Specialist agents (parallel layer)", 2)
    add_comparison_table(
        doc,
        ["#", "Agent", "Responsibility", "Typical data / APIs"],
        [
            ("3", "Weather Agent", "Real-time and seasonal weather guidance", "Open-Meteo geocoding, forecast, archive"),
            ("4", "Hotel Agent", "Best hotels by location, budget, and rating", "TripAdvisor16 via RapidAPI"),
            ("5", "Restaurants Agent", "Top restaurants and local food spots", "TripAdvisor16 via RapidAPI"),
            ("6", "Attractions Agent", "Must-visit places, landmarks, activities", "OpenTripMap radius search"),
            ("7", "Itinerary Agent", "Day-wise itinerary with logical flow", "LLM + outputs from weather, hotels, food, POIs"),
            ("8", "Budget Agent", "Total cost, daily budget, tiered breakdown", "User budget + FX API + listing context"),
            ("9", "Things to Carry Agent", "Packing list by destination and season", "Weather Agent output + destination metadata"),
        ],
    )

    add_heading(doc, "9.3 Merge and Summary Plan", 2)
    add_body(
        doc,
        "After parallel execution, the Merge node collects JSON/HTML/markdown fragments from each branch. "
        "The Summary Plan node (final green stage in Figure 3) synthesizes a single user-facing artifact: "
        "weather prose, hotel and restaurant cards, attraction suggestions, itinerary tables, budget tiers, "
        "and packing recommendations—stored as travel_details in Supabase and rendered on the dashboard.",
    )

    add_heading(doc, "9.4 FastAPI, LangGraph, and Travel Chatbot", 2)
    add_body(
        doc,
        "FastAPI is the HTTP gateway: it validates PlanRequest, invokes the LangGraph graph for /plan, and "
        "may expose dedicated routes (/itinerary, /budget) that re-run or refresh specific graph nodes. "
        "The Travel Chatbot (POST /chat) is auxiliary: it answers follow-up questions using the saved Summary "
        "Plan as context and maintains per-email session history—it is not part of the parallel fan-out in "
        "Figure 3 but complements the merged plan after generation.",
    )

    add_heading(doc, "9.5 Example user journey", 2)
    add_numbered_fresh(
        doc,
        [
            "User submits: “Plan a 5-day trip to Jaipur for 2 adults, budget ₹50,000, balanced style.”",
            "Parser Agent extracts Jaipur, dates, party size, budget, and style; triggers parallel agents.",
            "Weather, Hotel, Restaurants, Attractions, Itinerary, Budget, and Things to Carry agents run together.",
            "Merge collects all branch outputs; Summary Plan builds the dashboard trip brief.",
            "Plan is saved to Supabase; user reviews cards and tables on the dashboard.",
            "User asks Travel Chatbot: “Which hotel is best for families?” using stored plan context.",
        ],
    )

    add_figure(doc, ASSETS / "fig4_ui_mockup.png", "Figure 4: Conceptual dashboard layout for Go Bharat")

    add_heading(doc, "9.6 Prompt and tool notes per agent", 2)
    add_bullets(
        doc,
        [
            "Parser Agent: entity extraction and routing instructions; outputs structured fields for downstream nodes.",
            "Weather Agent: two short prose paragraphs; no bullet spam; uses current + seasonal snapshots.",
            "Hotel / Restaurants Agents: grounded on API listings; no invented venue names.",
            "Attractions Agent: categories and distances from OpenTripMap.",
            "Itinerary Agent: markdown day tables; 12-hour start times; Title Case day labels.",
            "Budget Agent: eight sections, three tiers (~35%, ~60%, 100% of user budget).",
            "Things to Carry Agent: packing list conditioned on weather and trip type.",
            "Travel Chatbot: system prompt embeds full merged plan from Supabase.",
        ],
    )

    doc.add_page_break()

    # ===== 10-18 remaining sections =====
    add_heading(doc, "10. Technologies Used", 1)
    tech_rows = [
        ("Python 3.12", "Language for services and agents."),
        ("FastAPI", "REST API and application entry (not Flask)."),
        ("Uvicorn", "ASGI server."),
        ("Pydantic / pydantic-settings", "Validation and configuration."),
        ("Jinja2", "Server-side HTML templates."),
        ("HTML / CSS", "UI structure and styling."),
        ("LangGraph", "Multi-agent graph: Parser, parallel specialists, Merge, Summary Plan."),
        ("LangChain + langchain-groq", "Prompt chains and Groq LLM inside each graph node."),
        ("httpx", "Async HTTP for external APIs."),
        ("Supabase", "PostgreSQL-backed plan storage."),
        ("Open-Meteo", "Geocoding and weather/archive."),
        ("RapidAPI TripAdvisor16", "Hotels and restaurants."),
        ("OpenTripMap", "Attractions near coordinates."),
        ("Exchange Rate API", "Budget conversion narrative."),
        ("Render", "Cloud hosting."),
        ("Groq API", "LLM inference."),
    ]
    add_comparison_table(doc, ["Technology", "Role"], tech_rows)

    from report_appendix import add_database_chapter

    add_database_chapter(doc, ASSETS)

    add_heading(doc, "12. Methodology", 1)
    phases = [
        "Phase 1 — Requirements: interviews within the team, study of tourism pain points, definition of MVP features.",
        "Phase 2 — Architecture: chose FastAPI + Supabase + multi-agent LLM layout; drew API and agent diagrams.",
        "Phase 3 — API integration: implemented PlannerService with httpx and error handling for missing keys.",
        "Phase 4 — Agent development: iterative prompt tuning for weather, itinerary, budget, and chat.",
        "Phase 5 — Frontend: built login/dashboard templates and styled result cards.",
        "Phase 6 — Persistence: Supabase insert/select by email_id and plan id.",
        "Phase 7 — Testing: manual tests with sample cities (e.g., Jaipur, Paris) and edge cases (invalid dates).",
        "Phase 8 — Deployment: Render environment variables and public URL for submission.",
    ]
    add_numbered_fresh(doc, phases)

    add_heading(doc, "13. Testing Approach", 1)
    add_bullets(
        doc,
        [
            "Health endpoint GET /api/v1/health for uptime checks.",
            "Validation test: check_out before check_in returns HTTP 422.",
            "Plan generation with and without optional API keys (graceful degradation messages).",
            "Supabase round-trip: store plan, list plans, fetch by id.",
            "Itinerary and budget endpoints after a plan exists.",
            "Chat multi-turn: verify session history per email_id in memory.",
            "UI walkthrough on desktop browser for form submission and results display.",
        ],
    )

    add_heading(doc, "14. Advantages and Limitations", 1)
    add_heading(doc, "14.1 Advantages", 2)
    add_bullets(
        doc,
        [
            "Unified platform reducing tab-switching across travel websites.",
            "Live data improves trust compared to purely generative answers.",
            "Multi-agent design improves maintainability and academic explainability.",
            "Modular features (plan → itinerary → budget → chat) match real user workflows.",
            "Deployed instance demonstrates full-stack capability beyond localhost.",
        ],
    )
    add_heading(doc, "14.2 Limitations", 2)
    add_bullets(
        doc,
        [
            "Email-based identification is not full authentication.",
            "API keys required for hotels/restaurants/attractions; missing keys reduce richness.",
            "LLM outputs may still need human verification for critical bookings.",
            "Chat session history is in-process memory, not persisted in Supabase.",
            "Hidden-gem discovery depends on OpenTripMap coverage, not curated local guides.",
        ],
    )

    add_heading(doc, "15. API Overview", 1)
    add_bullets(
        doc,
        [
            "GET /api/v1/health",
            "POST /api/v1/plan",
            "GET /api/v1/plans/{email_id}",
            "GET /api/v1/plans/{email_id}/{plan_id}",
            "POST /api/v1/chat",
            "GET /api/v1/chat/history/{email_id}",
            "POST /api/v1/chat/clear/{email_id}",
            "POST /api/v1/itinerary/{email_id}",
            "POST /api/v1/budget/{email_id}",
        ],
    )

    add_heading(doc, "16. Conclusion", 1)
    for para in [
        "Go Bharat shows how a student team can deliver a credible capstone at the intersection of web engineering, "
        "API integration, and agentic AI. The project is not merely a chatbot wrapper—it is an orchestrated system "
        "where live travel data and specialized agents cooperate.",
        "Choosing multiple agents over one monolithic model was a deliberate engineering decision. It improved output "
        "structure, reduced unnecessary recomputation, and made the codebase align with the diagrams and vocabulary "
        "expected in modern AI coursework (agents, orchestrator, tools, context).",
        "For Digital Bharat, tools that lower the friction of trip planning support domestic tourism and confidence "
        "among first-time travellers. We believe this architecture can evolve with authentication, booking partners, "
        "and richer local content in future semesters.",
    ]:
        add_body(doc, para)

    add_heading(doc, "17. Future Work", 1)
    add_bullets(
        doc,
        [
            "OAuth login and user-owned plan libraries.",
            "Persist chat history in Supabase for continuity across devices.",
            "Curated hidden-gem datasets for Indian cities.",
            "AQI, transit, and event APIs integrated as additional specialist agents.",
            "PDF/ICS export for itineraries.",
            "Hindi and regional-language prompts.",
            "Voice interface via speech-to-text agent.",
            "Automated unit and integration tests in CI.",
            "Rate limiting and cost monitoring for Groq usage.",
        ],
    )

    from report_appendix import add_appendices, add_extra_main_sections

    add_extra_main_sections(doc, ASSETS)
    add_appendices(doc)

    add_heading(doc, "18. References", 1)
    refs = [
        "National Digital Tourism Mission — https://tourism.gov.in/",
        "Ministry of Tourism annual reports — https://tourism.gov.in/media/annual-reports",
        "NITI Aayog Tourism — https://niti.gov.in/",
        "IBEF Tourism & Hospitality — https://www.ibef.org/industry/tourism-hospitality-india",
        "FastAPI — https://fastapi.tiangolo.com/",
        "LangChain — https://python.langchain.com/",
        "Groq — https://console.groq.com/",
        "Supabase — https://supabase.com/docs",
        "Open-Meteo — https://open-meteo.com/",
        "Render — https://render.com/docs",
        "RapidAPI — https://rapidapi.com/",
        "Multi-agent systems (conceptual) — Russell & Norvig, Artificial Intelligence: A Modern Approach",
        f"Live deployment — {LIVE_URL}",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref, style="List Bullet")
        style_paragraph(p, size=10)

    return doc


def main() -> None:
    doc = build_document()
    doc.save(str(OUT_PATH))
    print(f"Saved: {OUT_PATH}")
    from docx import Document as _D

    d = _D(str(OUT_PATH))
    words = sum(len(p.text.split()) for p in d.paragraphs)
    print(f"Words: ~{words} | Figures: 10 | Tables: {len(d.tables)}")


if __name__ == "__main__":
    main()
