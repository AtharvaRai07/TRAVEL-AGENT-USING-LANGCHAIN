"""Extended sections and database deep-dive for the Capstone report."""
from __future__ import annotations

from pathlib import Path

from docx import Document

from generate_capstone_report import (
    LIVE_URL,
    add_body,
    add_bullets,
    add_comparison_table,
    add_figure,
    add_heading,
    add_horizontal_rule,
    add_numbered,
)


def add_database_chapter(doc: Document, assets: Path) -> None:
    add_heading(doc, "11. Database Design (Detailed)", 1)
    add_body(
        doc,
        "Persistence is a core requirement for Go Bharat: users must return to the dashboard, list past trips, "
        "and invoke itinerary, budget, and chat endpoints without re-entering data. We chose Supabase because it "
        "provides managed PostgreSQL, a simple REST/JS client, and fast setup for academic projects. The application "
        "does not normalize hotels or weather into separate tables; instead it stores one rich JSON document per plan, "
        "which matches how the frontend and agents consume data.",
    )
    add_body(
        doc,
        "Figure 5 presents the logical database layout. The physical implementation is a single table "
        "travel_plans in the public schema. Column travel_details is JSONB so PostgreSQL can index keys if we "
        "add GIN indexes later.",
    )
    add_figure(doc, assets / "fig5_database_er.png", "Figure 5: Entity layout — travel_plans table and JSONB structure")

    add_heading(doc, "11.1 Column definitions", 2)
    add_comparison_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ("id", "BIGINT (identity)", "Surrogate primary key; returned after INSERT; used in GET by plan id."),
            ("email_id", "TEXT", "User identifier from the form (validated as email in Pydantic). Groups all plans for one person."),
            ("city", "TEXT", "Destination string used for filtering when multiple plans exist."),
            ("travel_details", "JSONB", "Full plan payload: API results, HTML snippets, metadata, budget fields."),
            ("created_at", "TIMESTAMPTZ", "Server timestamp when the row was inserted; supports ordering by recency."),
        ],
    )

    add_heading(doc, "11.2 CRUD operations in the application", 2)
    add_body(
        doc,
        "SupabaseService wraps four access patterns used by routes.py. Figure 6 maps each Python method to HTTP endpoints.",
    )
    add_figure(doc, assets / "fig6_database_crud.png", "Figure 6: Database operations performed by SupabaseService")

    add_body(doc, "store_travel_plan(email_id, city, travel_details) executes INSERT after POST /plan succeeds.")
    add_body(
        doc,
        "fetch_travel_plan(email_id, city=None) SELECTs rows matching email_id, optionally city, ORDER BY id DESC "
        "LIMIT 1, then merges the JSON into the top-level dict returned to agents.",
    )
    add_body(doc, "fetch_all_plans(email_id) returns a lightweight list for the dashboard (id, city, destination, dates).")
    add_body(doc, "fetch_plan_by_id(email_id, plan_id) fetches one row for detail views.")

    add_heading(doc, "11.3 Sample record", 2)
    add_figure(
        doc,
        assets / "fig7_sample_record.png",
        "Figure 7: Sample travel_plans row — column and JSON field values",
        width=6.5,
    )

    add_heading(doc, "11.4 SQL script for Supabase", 2)
    add_body(doc, "The following script can be run in the Supabase SQL Editor to create the schema:")
    add_body(
        doc,
        "-- Go Bharat schema\n"
        "CREATE TABLE IF NOT EXISTS public.travel_plans (\n"
        "  id BIGINT GENERATED ALWAYS AS IDENTITY PRIMARY KEY,\n"
        "  email_id TEXT NOT NULL,\n"
        "  city TEXT NOT NULL,\n"
        "  travel_details JSONB NOT NULL DEFAULT '{}'::jsonb,\n"
        "  created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()\n"
        ");\n"
        "CREATE INDEX IF NOT EXISTS idx_travel_plans_email_id\n"
        "  ON public.travel_plans (email_id);\n"
        "CREATE INDEX IF NOT EXISTS idx_travel_plans_email_id_city\n"
        "  ON public.travel_plans (email_id, city);\n"
        "CREATE INDEX IF NOT EXISTS idx_travel_plans_travel_details_gin\n"
        "  ON public.travel_plans USING GIN (travel_details);",
    )

    add_heading(doc, "11.5 Design rationale: JSONB vs many tables", 2)
    for para in [
        "A normalized design might split hotels, restaurants, and weather into child tables. For Capstone-I scope, "
        "that adds join complexity without changing the UI, which renders pre-built HTML strings. JSONB keeps one "
        "round-trip per plan fetch and mirrors the PlanResponse Pydantic model after model_dump().",
        "Trade-off: analytics across all users' favourite cities requires JSON operators (e.g. travel_details->>'destination'). "
        "For future work, materialized views or ETL could extract dimensions from JSONB.",
        "Security note: the backend should use the Supabase service role key only on the server. Row Level Security "
        "can be enabled for production if the anon key is ever exposed to browsers.",
    ]:
        add_body(doc, para)

    doc.add_page_break()


def add_extra_main_sections(doc: Document, assets: Path) -> None:
    add_heading(doc, "19. Deployment Architecture", 1)
    add_body(
        doc,
        "The production instance is hosted on Render as a Python web service (see render.yaml). Environment variables "
        "inject SUPABASE_URL, SUPABASE_KEY, GROQ_API_KEY, RAPIDAPI_KEY, OPENTRIP_API_KEY, and related settings at runtime. "
        "Figure 9 shows how the browser, Render container, Supabase, Groq, and external travel APIs interact.",
    )
    add_figure(doc, assets / "fig9_deployment.png", "Figure 9: Deployment on Render with Supabase and external APIs")

    add_heading(doc, "20. Layered Software Architecture", 1)
    add_body(
        doc,
        "The codebase follows a conventional layered structure so coursework evaluators can map files to responsibilities. "
        "Figure 10 summarizes the layers from templates down to external I/O.",
    )
    add_figure(doc, assets / "fig10_layered_architecture.png", "Figure 10: Layered architecture of the Go Bharat codebase")

    add_body(doc, "Key directories:")
    add_bullets(
        doc,
        [
            "app/main.py — FastAPI app, static mount, page routes for / and /dashboard.",
            "app/api/routes.py — REST API under /api/v1.",
            "app/services/planner.py — Async orchestration of travel APIs.",
            "app/services/*_agent.py — LangChain specialist agents.",
            "app/services/supabase_service.py — Database access.",
            "app/schemas/travel.py — Pydantic request/response models.",
            "app/templates/ — Jinja2 HTML.",
            "app/static/css/ — Stylesheets.",
        ],
    )

    add_heading(doc, "21. API Request Sequence", 1)
    add_body(doc, "Figure 8 traces the main happy-path when a user creates a plan.")
    add_figure(doc, assets / "fig8_api_sequence.png", "Figure 8: Sequence diagram for POST /api/v1/plan")

    doc.add_page_break()


def add_appendices(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "Appendix A — Sample API Request (POST /api/v1/plan)", 1)
    add_body(doc, "Content-Type: application/json")
    add_body(
        doc,
        '{\n'
        '  "email_id": "student@example.com",\n'
        '  "city": "Jaipur",\n'
        '  "check_in": "2026-04-10",\n'
        '  "check_out": "2026-04-14",\n'
        '  "adults": 2,\n'
        '  "budget_amount": 50000,\n'
        '  "budget_currency": "INR",\n'
        '  "style": "balanced"\n'
        "}",
    )
    add_body(
        doc,
        "The server validates types and constraints (e.g. check_out must be after check_in). On success, the response "
        "includes weather, hotels, restaurants, attractions, currency, final_response HTML, and generated_at. "
        "A copy is inserted into Supabase unless storage fails—in which case a warning field may appear on the response.",
    )

    add_heading(doc, "Appendix B — Environment Variables", 1)
    add_comparison_table(
        doc,
        ["Variable", "Purpose"],
        [
            ("SUPABASE_URL", "Supabase project API URL."),
            ("SUPABASE_KEY", "Service or anon key for database client."),
            ("GROQ_API_KEY", "Authentication for ChatGroq LLM calls."),
            ("RAPIDAPI_KEY", "TripAdvisor16 hotel and restaurant APIs."),
            ("OPENTRIP_API_KEY", "OpenTripMap attraction search."),
            ("ENVIRONMENT / DEBUG", "Render production settings."),
        ],
    )

    add_heading(doc, "Appendix C — Tourism Context in India", 1)
    for para in [
        "India's tourism sector contributes significantly to GDP and employment. Domestic travel has grown with "
        "improved connectivity and digital payments. Government initiatives emphasize digital discovery of destinations, "
        "safety, and sustainable tourism.",
        "Go Bharat aligns with this direction by lowering the effort to assemble a first-pass itinerary for any city, "
        "whether a heritage centre like Jaipur or an international destination entered by the user. The project is "
        "educational: it demonstrates integration skills rather than replacing licensed tour operators.",
        "Market platforms (OTAs, maps, review sites) remain authoritative for bookings. Our site aggregates information "
        "and adds narrative AI layers; users should verify prices and availability before paying.",
    ]:
        add_body(doc, para)

    add_heading(doc, "Appendix D — Glossary", 1)
    glossary = [
        ("Agentic AI", "AI components with goals, context, and structured outputs, orchestrated by code."),
        ("LLM", "Large Language Model; here accessed via Groq."),
        ("JSONB", "PostgreSQL binary JSON type with indexing support."),
        ("Orchestrator", "FastAPI + PlannerService; decides which agent or API to call."),
        ("RAG-like context", "Injecting stored plan text into prompts without a vector database."),
        ("ASGI", "Asynchronous Server Gateway Interface; served by Uvicorn."),
    ]
    add_comparison_table(doc, ["Term", "Meaning"], glossary)

    add_heading(doc, "Appendix E — Extended Single vs Multi-Agent Discussion", 1)
    for para in [
        "In industry, 'single agent' products sometimes mean one chat thread with tool calling (function calling). "
        "Even then, tools are separate functions—the model chooses among them. Go Bharat makes this explicit in code: "
        "the Weather Agent is not the same node as the Budget Agent in the LangGraph, and each has a maintained prompt "
        "(embedded in the class constructor).",
        "Failure isolation example: if the budget endpoint times out, the user still has the base plan and weather "
        "brief in Supabase. A monolithic agent would require regenerating everything or manual editing of a long response.",
        "Token economics: a budget table prompt may use max_tokens=2000 while weather uses 700. Tunings would conflict "
        "in a single shared chain.",
        "Teaching value: evaluators can assign marks for architecture diagrams that match the repository—an important "
        "goal for Capstone-I documentation.",
    ]:
        add_body(doc, para)

    add_heading(doc, "Appendix F — External API Integrations", 1)
    apis = [
        (
            "Open-Meteo Geocoding",
            "Resolves city name to latitude, longitude, and display name. Scoring logic in PlannerService "
            "picks the best match when multiple places share a name.",
        ),
        (
            "Open-Meteo Forecast & Archive",
            "Current temperature and humidity for 'right now'; historical archive for the same calendar date "
            "one year prior to estimate seasonal highs, lows, and rain.",
        ),
        (
            "RapidAPI TripAdvisor16",
            "searchLocation + searchHotels for accommodation; parallel flow for restaurants. Requires RAPIDAPI_KEY.",
        ),
        (
            "OpenTripMap",
            "Radius search for tourist attractions near coordinates; fetches xid details for names and categories.",
        ),
        (
            "Exchange Rate API",
            "Converts user budget from chosen currency to a reference currency for explanatory text.",
        ),
        (
            "Groq ChatGroq",
            "Hosts the openai/gpt-oss-120b model used by all LangChain chains.",
        ),
    ]
    for name, desc in apis:
        add_heading(doc, name, 3)
        add_body(doc, desc)

    add_heading(doc, "Appendix G — Live Application", 1)
    add_body(
        doc,
        f"The deployed build is available at {LIVE_URL}. Evaluators can test the dashboard, generate a plan, and "
        "optionally exercise chat, itinerary, and budget endpoints via the UI or API client.",
    )
    add_body(
        doc,
        "Recommended demo script for viva: (1) open dashboard, (2) create plan for Jaipur with INR budget, "
        "(3) show weather and hotel cards, (4) generate itinerary, (5) generate budget tiers, (6) ask chatbot "
        "one question about restaurants, (7) show Supabase row in Table Editor if projector allows.",
    )

    add_heading(doc, "Appendix H — Risk and Mitigation", 1)
    add_comparison_table(
        doc,
        ["Risk", "Mitigation"],
        [
            ("API key leakage", "Keys only in Render env vars / .env; never committed to Git."),
            ("LLM hallucination", "Agents receive API-sourced lists; prompts ask to cite provided venues."),
            ("Supabase outage", "Plan still returns; warning field if storage fails."),
            ("Rate limits on Groq/RapidAPI", "On-demand agents; user triggers heavy calls separately."),
            ("Incorrect dates", "Pydantic + route returns HTTP 422 if check_out ≤ check_in."),
        ],
    )

    add_horizontal_rule(doc)


def extend_report(doc: Document, assets: Path) -> None:
    add_database_chapter(doc, assets)
    add_extra_main_sections(doc, assets)
    add_appendices(doc)
